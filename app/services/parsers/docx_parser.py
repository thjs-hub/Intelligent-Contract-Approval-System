"""DOCX 合同文档解析器 (M04)。

使用 python-docx 库读取 DOCX 文件全文，然后调用通用文本提取器进行字段提取。
"""

from typing import Any

from app.services.parsers.base import get_extractor


class DocxParser:
    """DOCX 合同文档解析器 — 基础版（第二阶段）

    解析流程:
      1. 使用 python-docx 读取所有段落文本
      2. 拼接为全文
      3. 调用 RegexTextExtractor 提取基本信息和条款信息
    """

    def parse(self, file_path: str) -> dict[str, Any]:
        """解析 DOCX 文件

        参数:
          file_path: DOCX 文件绝对路径

        返回:
          {
            "basic_info_json": dict,
            "clause_info_json": dict,
          }
        """
        from docx import Document

        # 1. 读取 DOCX 文件全文
        doc = Document(file_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # 2. 调用通用文本提取器
        extractor = get_extractor()
        basic_info = extractor.extract_basic_info(full_text)
        clause_info = extractor.extract_clauses(full_text)

        return {
            "basic_info_json": basic_info,
            "clause_info_json": clause_info,
        }

    def extract_text(self, file_path: str) -> str:
        """仅提取 DOCX 全文文本（不含结构化字段）

        供 M05 OCR 完成后复用 M04 的字段提取逻辑使用。
        """
        from docx import Document

        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
