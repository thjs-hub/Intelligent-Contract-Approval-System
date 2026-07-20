"""Mock 审批系统适配器。

用于本地开发和单元测试，模拟企业审批系统的全部接口行为:
  - 拉取待办列表: 返回固定模式的测试审批单
  - 查询详情: 返回包含附件列表的详情
  - 下载附件: 返回模拟的文件内容
  - 写入评论: 模拟成功响应
"""

from typing import Any

from app.services.adapters.base import ApprovalSystemAdapter


class MockApprovalAdapter(ApprovalSystemAdapter):
    """Mock 审批系统适配器

    所有接口返回预设的模拟数据，便于在无真实审批系统环境下完成端到端联调。
    """

    # 预置的审批单模板（按索引生成）
    _APPROVAL_TEMPLATE = {
        "applicant_name": "张三",
        "form_data": {"合同金额": "100000", "币种": "CNY"},
    }

    # 预置附件清单（每个审批单 2 个附件）
    _ATTACHMENTS = [
        {"attachment_id": "att-001", "file_name": "采购合同.docx", "file_type": "docx"},
        {"attachment_id": "att-002", "file_name": "附件协议.pdf", "file_type": "pdf"},
    ]

    async def fetch_pending_approvals(self, limit: int = 20) -> list[dict[str, Any]]:
        """返回模拟的待处理审批单列表

        生成 limit 条以 AP-2026-XXXX 为编号的审批单数据。
        """
        return [
            {
                "approval_code": f"AP-2026-{i:04d}",
                "approval_title": f"测试合同审批单-{i}",
                "applicant_name": self._APPROVAL_TEMPLATE["applicant_name"],
                "attachment_count": len(self._ATTACHMENTS),
            }
            for i in range(1, limit + 1)
        ]

    async def fetch_approval_detail(self, instance_id: str) -> dict[str, Any]:
        """返回模拟的审批单详情"""
        return {
            "approval_code": instance_id,
            "approval_title": f"合同审批-{instance_id}",
            "applicant_name": self._APPROVAL_TEMPLATE["applicant_name"],
            "form_data": self._APPROVAL_TEMPLATE["form_data"],
            "attachments": self._ATTACHMENTS.copy(),
            "status": "pending",
        }

    async def download_attachment(self, attachment_id: str) -> bytes:
        """返回模拟的附件二进制内容

        根据附件 ID 生成真实的 DOCX 文件内容，便于测试 M04 文档解析流程。
        """
        from io import BytesIO

        from docx import Document

        # 创建真实的 DOCX 文档（python-docx 可解析）
        doc = Document()
        doc.add_paragraph("采购合同")
        doc.add_paragraph(f"合同编号：MOCK-{attachment_id}")
        doc.add_paragraph("甲方：甲方科技有限公司")
        doc.add_paragraph("乙方：乙方贸易有限公司")
        doc.add_paragraph("合同金额：人民币壹拾万元整（¥100,000）")
        doc.add_paragraph("本合同自2026年1月1日起生效，有效期至2026年12月31日。")
        doc.add_paragraph("")
        doc.add_paragraph("付款方式：合同签订后7日内支付预付款30%，货到验收后支付60%，质保期满后支付10%。")
        doc.add_paragraph("交付时间：卖方应于合同签订后30日内完成交付。")
        doc.add_paragraph("验收标准：按合同附件技术规格书执行。")
        doc.add_paragraph("违约责任：任一方违约应支付合同金额10%的违约金。")
        doc.add_paragraph("保密条款：双方对合同内容负有保密义务，保密期限为合同终止后3年。")
        doc.add_paragraph("争议解决：本合同争议提交合同签订地有管辖权的人民法院诉讼解决。")

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    async def write_comment(self, instance_id: str, comment_text: str) -> dict[str, Any]:
        """模拟写入评论成功"""
        return {
            "success": True,
            "instance_id": instance_id,
            "comment_id": f"comment-{instance_id}-{hash(comment_text) & 0xFFFFFFFF}",
            "message": "评论写入成功（Mock）",
        }
