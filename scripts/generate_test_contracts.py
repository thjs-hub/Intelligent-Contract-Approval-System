"""Generate two test contracts (.docx) with deliberate minor errors for review testing."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

UPLOADS_DIR = os.path.join(os.path.dirname(__file__), "..", "uploads")


def set_cell_text(cell, text, bold=False, size=11):
    """Set cell paragraph text with styling."""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.name = "SimSun"
    run.bold = bold


def add_heading_styled(doc, text, level=1):
    """Add a heading with centering."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "SimHei"
    return heading


def add_clause(doc, title, content):
    """Add a numbered clause."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.name = "SimSun"
    run_content = p.add_run(content)
    run_content.font.size = Pt(11)
    run_content.font.name = "SimSun"


# ============================================================
# Contract 1: 产品采购合同
# Errors: 预付款50%(R001), 验收标准模糊(R011), 缺知识产权(R010),
#          管辖地不明确(R005), 缺数据隐私条款(R009)
# ============================================================
def create_contract_1():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("产品采购合同")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "SimHei"

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("合同编号：CG-2026-0718-001")
    run.font.size = Pt(10)
    run.font.name = "SimSun"

    doc.add_paragraph()  # blank line

    # --- Parties ---
    doc.add_heading("一、合同双方", level=2)
    doc.add_paragraph(
        "甲方（采购方）：深圳星河科技有限公司\n"
        "统一社会信用代码：91440300MA5DXXXX1X\n"
        "法定代表人：张明远\n"
        "地址：深圳市南山区科技园南路18号星河科技大厦12层\n"
        "联系人：李经理　　电话：0755-6688XXXX"
    )
    doc.add_paragraph(
        "乙方（供货方）：北京华威电子有限公司\n"
        "统一社会信用代码：91110108MA0XXXXX2Y\n"
        "法定代表人：王建国\n"
        "地址：北京市海淀区中关村北大街66号中科产业园B座5层\n"
        "联系人：赵总监　　电话：010-8288XXXX"
    )

    # --- Subject ---
    doc.add_heading("二、合同标的", level=2)
    doc.add_paragraph(
        "甲方向乙方采购以下产品：\n"
        "1. 高性能服务器主板，型号 HW-S1000，数量 200 台，单价 8,500 元；\n"
        "2. 企业级固态硬盘，型号 HW-SSD-2T，数量 500 块，单价 1,200 元；\n"
        "3. 万兆网络交换机，型号 HW-SW-10G，数量 50 台，单价 6,800 元。"
    )

    # --- Amount ---
    doc.add_heading("三、合同金额", level=2)
    doc.add_paragraph(
        "合同总金额：人民币贰佰陆拾肆万元整（¥2,640,000.00）\n"
        "币种：人民币\n"
        "上述金额含增值税（税率 13%），包含运输费、保险费及安装调试费。"
    )

    # --- Payment (ERROR: prepayment 50%) ---
    doc.add_heading("四、付款方式", level=2)
    add_clause(
        doc,
        "4.1 预付款：",
        "本合同签订后 7 个工作日内，甲方向乙方支付合同总金额的 50% 作为预付款，"
        "计人民币壹佰叁拾贰万元整（¥1,320,000.00）。"
        "乙方收到预付款后安排生产。"
    )
    add_clause(
        doc,
        "4.2 到货款：",
        "全部货物运抵甲方指定地点并完成初步验收后 10 个工作日内，"
        "甲方向乙方支付合同总金额的 40%，计人民币壹佰零伍万陆仟元整（¥1,056,000.00）。"
    )
    add_clause(
        doc,
        "4.3 质保金：",
        "剩余 10% 作为质量保证金，质保期满后 15 个工作日内支付，计人民币贰拾陆万肆仟元整（¥264,000.00）。"
        "质保期为验收合格之日起 12 个月。"
    )

    # --- Delivery ---
    doc.add_heading("五、交付条款", level=2)
    doc.add_paragraph(
        "交付时间：乙方应于合同签订之日起 45 个工作日内完成全部货物的交付。\n"
        "交付地点：深圳市南山区科技园南路18号星河科技大厦12层，甲方指定仓库。\n"
        "运输方式：由乙方负责安排物流运输，费用由乙方承担。\n"
        "货物风险自甲方签收之时起转移至甲方。"
    )

    # --- Acceptance (ERROR: vague standard) ---
    doc.add_heading("六、验收标准", level=2)
    doc.add_paragraph(
        "验收标准：按行业标准执行。乙方提供产品合格证及出厂检测报告，"
        "甲方在收到货物后 5 个工作日内完成验收。如甲方逾期未提出书面异议，视为验收合格。"
    )

    # --- Breach ---
    doc.add_heading("七、违约责任", level=2)
    add_clause(
        doc,
        "7.1 ",
        "如乙方未能按时交付货物，每逾期一日，应向甲方支付未交付部分货值 0.05% 的违约金，"
        "违约金总额不超过合同总金额的 10%。"
    )
    add_clause(
        doc,
        "7.2 ",
        "如甲方未按约定时间支付款项，每逾期一日，应向乙方支付逾期付款金额 0.05% 的违约金。"
    )
    add_clause(
        doc,
        "7.3 ",
        "如交付产品存在质量问题，乙方应在收到甲方书面通知后 10 个工作日内负责免费更换或维修。"
    )

    # --- Warranty ---
    doc.add_heading("八、保修条款", level=2)
    doc.add_paragraph(
        "乙方对所提供的产品提供自验收合格之日起 36 个月的免费保修服务。"
        "保修期内，因产品自身质量问题导致的故障，乙方负责免费维修或更换。"
        "人为损坏、不可抗力因素导致的损坏不在保修范围内。"
    )

    # --- Confidentiality ---
    doc.add_heading("九、保密条款", level=2)
    doc.add_paragraph(
        "双方对在合同签订及履行过程中知悉的对方商业秘密、技术信息及其他未公开信息"
        "负有保密义务。未经对方书面同意，任何一方不得向第三方泄露。"
        "本保密条款在合同终止后 3 年内持续有效。"
    )

    # NOTE: Missing clause 十 - 数据隐私条款 (trigger R009)
    # NOTE: Missing clause 十一 - 知识产权条款 (trigger R010)

    # --- Dispute (ERROR: vague jurisdiction) ---
    doc.add_heading("十、争议解决", level=2)
    doc.add_paragraph(
        "因本合同引起的或与本合同有关的任何争议，双方应首先友好协商解决；"
        "协商不成的，任何一方均有权向合同签订地有管辖权的人民法院提起诉讼。"
    )

    # --- Miscellaneous ---
    doc.add_heading("十一、其他条款", level=2)
    doc.add_paragraph(
        "本合同自双方签字盖章之日起生效，有效期自 2026 年 7 月 21 日至 2027 年 7 月 20 日。\n"
        "本合同一式肆份，甲乙双方各执贰份，具有同等法律效力。\n"
        "本合同附件为本合同不可分割的组成部分，与本合同具有同等法律效力。"
    )

    # Signatures
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig.add_run("甲方（盖章）：深圳星河科技有限公司").font.size = Pt(11)
    sig.add_run("　　　　　　　　乙方（盖章）：北京华威电子有限公司").font.size = Pt(11)
    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig2.add_run("授权代表签字：_______________").font.size = Pt(11)
    sig2.add_run("　　　　　　授权代表签字：_______________").font.size = Pt(11)
    sig3 = doc.add_paragraph()
    sig3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig3.add_run("日期：2026 年    月    日").font.size = Pt(11)
    sig3.add_run("　　　　　　　　　日期：2026 年    月    日").font.size = Pt(11)

    path = os.path.join(UPLOADS_DIR, "采购合同_CG-2026-0718-001.docx")
    os.makedirs(UPLOADS_DIR, exist_ok=True)
    doc.save(path)
    print(f"[OK] Created: {path}")
    return path


# ============================================================
# Contract 2: 技术服务合同
# Errors: 金额模糊(R007), 违约责任不对等(R004), 缺保密(R008),
#          自动续约(R003), 甲方信息不全(R006)
# ============================================================
def create_contract_2():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("技术服务合同")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "SimHei"

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("合同编号：JS-2026-0721-002")
    run.font.size = Pt(10)
    run.font.name = "SimSun"

    doc.add_paragraph()

    # --- Parties (ERROR: Party A lacks address detail) ---
    doc.add_heading("一、合同双方", level=2)
    doc.add_paragraph(
        "甲方（委托方）：上海云帆数据有限公司\n"
        "法定代表人：陈海峰\n"
        "联系人：王经理　　电话：021-5388XXXX\n"
        "（甲方信息详见附件一）"
    )
    doc.add_paragraph(
        "乙方（服务方）：杭州智睿信息技术有限公司\n"
        "统一社会信用代码：91330100MA2XXXXX3Z\n"
        "法定代表人：刘文博\n"
        "地址：浙江省杭州市滨江区网商路 599 号网易大厦 A 座 15 层\n"
        "联系人：孙经理　　电话：0571-8899XXXX"
    )

    # --- Service scope ---
    doc.add_heading("二、服务内容", level=2)
    doc.add_paragraph(
        "乙方为甲方提供以下技术服务：\n"
        "1. 企业数据中台搭建与部署，包括数据采集、清洗、存储及可视化；\n"
        "2. 现有业务系统（ERP、CRM）数据迁移与整合；\n"
        "3. 数据安全策略规划与实施；\n"
        "4. 甲方技术团队培训（不少于 40 学时）。\n"
        "具体服务内容及技术规格详见附件二《技术需求说明书》。"
    )

    # --- Amount (ERROR: some amounts vague) ---
    doc.add_heading("三、合同金额及支付方式", level=2)
    doc.add_paragraph(
        "3.1 基础服务费：人民币壹佰捌拾万元整（¥1,800,000.00），包含数据中台搭建、"
        "系统迁移及培训费用。\n"
        "3.2 扩展服务费：如甲方需要超出附件二范围的额外功能开发，相关费用由双方另行协商确定。\n"
        "3.3 支付方式：\n"
        "　(1) 合同签订后 10 个工作日内支付基础服务费的 30%，计人民币伍拾肆万元整（¥540,000.00）；\n"
        "　(2) 数据中台上线运行后 10 个工作日内支付基础服务费的 40%，计人民币柒拾贰万元整（¥720,000.00）；\n"
        "　(3) 项目整体验收通过后 10 个工作日内支付基础服务费的 30%，计人民币伍拾肆万元整（¥540,000.00）。"
    )

    # --- Service period ---
    doc.add_heading("四、服务期限", level=2)
    doc.add_paragraph(
        "本合同服务期限自 2026 年 8 月 1 日起至 2027 年 7 月 31 日止，共计 12 个月。\n"
        "生效日期：2026 年 8 月 1 日。\n"
        "合同到期前 30 日内，如双方均未提出书面异议，本合同自动续约一年。"
    )

    # --- Acceptance ---
    doc.add_heading("五、验收标准", level=2)
    doc.add_paragraph(
        "5.1 乙方完成数据中台部署后，甲方应在 15 个工作日内按照附件二《技术需求说明书》"
        "进行功能验收。\n"
        "5.2 验收内容包括但不限于：数据采集准确性、数据处理性能（查询响应时间不超过 3 秒）、"
        "系统稳定性（连续运行 72 小时无故障）。\n"
        "5.3 甲方应在验收完成后出具书面验收报告。如验收未通过，乙方应在 15 个工作日内完成整改。"
    )

    # --- Breach (ERROR: only limits one party) ---
    doc.add_heading("六、违约责任", level=2)
    doc.add_paragraph(
        "6.1 如乙方未能按照合同约定的时间完成各项服务交付，每逾期一日，应向甲方支付"
        "合同总金额 0.1% 的违约金，违约金总额不超过合同总金额的 20%。\n"
        "6.2 如乙方提供的服务存在严重质量问题导致甲方无法正常使用，乙方应退还已收取的全部费用，"
        "并赔偿甲方因此遭受的直接经济损失。\n"
        "6.3 如乙方将本合同约定的服务转包或分包给第三方，甲方有权单方面解除合同，"
        "乙方应退还全部已付款项。"
    )

    # --- IP (has this one - good) ---
    doc.add_heading("七、知识产权", level=2)
    doc.add_paragraph(
        "7.1 乙方为履行本合同而开发的软件源代码、技术文档及相关成果的知识产权归甲方所有。\n"
        "7.2 乙方在执行本合同前已拥有的知识产权（包括但不限于基础框架、通用组件、"
        "算法模型等）仍归乙方所有，但乙方应向甲方提供永久、不可撤销的使用许可。\n"
        "7.3 任何一方不得在未经对方书面许可的情况下，将涉及对方知识产权的技术成果"
        "向第三方披露或授权使用。"
    )

    # --- Data / Privacy ---
    doc.add_heading("八、数据安全与隐私保护", level=2)
    doc.add_paragraph(
        "8.1 乙方在提供服务过程中接触到的甲方业务数据、客户信息及其他敏感数据，"
        "应严格遵循《中华人民共和国数据安全法》及《中华人民共和国个人信息保护法》的相关规定。\n"
        "8.2 乙方不得将甲方数据用于本合同约定之外的任何目的，不得复制、修改、出售或泄露甲方数据。\n"
        "8.3 乙方应采取符合行业标准的技术手段（包括但不限于数据加密、访问控制、日志审计）"
        "保障甲方数据安全。\n"
        "8.4 乙方应在合同终止后 30 日内彻底删除其持有的所有甲方数据，并出具书面删除证明。"
    )

    # NOTE: Missing clause - 保密条款 (trigger R008)

    # --- Dispute ---
    doc.add_heading("九、争议解决", level=2)
    doc.add_paragraph(
        "因本合同引起的或与本合同有关的任何争议，双方应友好协商解决；"
        "协商不成的，任何一方均有权向被告所在地有管辖权的人民法院提起诉讼。"
    )

    # --- Others ---
    doc.add_heading("十、其他", level=2)
    doc.add_paragraph(
        "本合同自双方签字盖章之日起生效。\n"
        "本合同一式肆份，甲乙双方各执贰份，具有同等法律效力。\n"
        "本合同附件包括：\n"
        "　附件一：甲方基本信息登记表\n"
        "　附件二：技术需求说明书"
    )

    # Signatures
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("甲方（盖章）：上海云帆数据有限公司").font.size = Pt(11)
    sig.add_run("　　　　　乙方（盖章）：杭州智睿信息技术有限公司").font.size = Pt(11)
    sig2 = doc.add_paragraph()
    sig2.add_run("授权代表签字：_______________").font.size = Pt(11)
    sig2.add_run("　　　　　授权代表签字：_______________").font.size = Pt(11)
    sig3 = doc.add_paragraph()
    sig3.add_run("日期：2026 年    月    日").font.size = Pt(11)
    sig3.add_run("　　　　　日期：2026 年    月    日").font.size = Pt(11)

    path = os.path.join(UPLOADS_DIR, "技术服务合同_JS-2026-0721-002.docx")
    os.makedirs(UPLOADS_DIR, exist_ok=True)
    doc.save(path)
    print(f"[OK] Created: {path}")
    return path


if __name__ == "__main__":
    create_contract_1()
    create_contract_2()
    print("\nDone. Two test contracts generated.")
